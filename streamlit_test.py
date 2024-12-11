import streamlit as st
import tempfile
from lxml import etree
import zipfile
from docx import Document
from utils_test import chunk_by_token, extract_text_from_pdf
from MyModel import ExtractA, ExtractB, ExtractC, ExtractD

# Lớp CustomDocumentConverter để xử lý file Word và PDF
class CustomDocumentConverter:
    def convert(self, file_path, file_type="docx"):
        if file_type == "docx":
            try:
                with zipfile.ZipFile(file_path, 'r') as z:
                    xml_content = z.read('word/document.xml')
                root = etree.XML(xml_content)
                namespaces = root.nsmap
                if "w" not in namespaces:
                    namespaces["w"] = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                elements = root.xpath(".//w:t", namespaces=namespaces)
                text = " ".join([el.text for el in elements if el.text])
                return text
            except KeyError:
                raise ValueError("File .docx không hợp lệ hoặc không chứa tài liệu Word!")
            except Exception as e:
                raise ValueError(f"Lỗi xử lý file Word: {e}")
        elif file_type == "pdf":
            return extract_text_from_pdf(file_path)
        else:
            raise ValueError("Định dạng file không được hỗ trợ!")

# Hàm xuất danh sách câu hỏi ra file Word
def export_to_word(questions, exam_title):
    doc = Document()
    doc.add_heading(exam_title, level=1)
    for idx, quiz in enumerate(questions, start=1):
        doc.add_heading(f"Câu {idx} ({quiz['level']}):", level=2)
        doc.add_paragraph(quiz["question"])
        for i, option in enumerate(quiz['choices'], start=65):  # 65 là mã ASCII của 'A'
            doc.add_paragraph(f"{chr(i)}. {option.strip()}")
        doc.add_paragraph(f"Đáp án đúng: {quiz['answer']}")
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

# Giao diện Streamlit
st.set_page_config(page_title="Website Tạo Đề Trắc Nghiệm Tự Động", layout="wide")

st.markdown(
    """
    <h1 style="text-align: center; color: #0056B3;">
        Website Tạo Đề Trắc Nghiệm Tự Động 🚀
    </h1>
    """,
    unsafe_allow_html=True
)

st.subheader("Tiêu đề đề kiểm tra:")
exam_title = st.text_input("", placeholder="Đề kiểm tra môn Tin học lớp 11")

st.subheader("Tải lên file nội dung (.docx hoặc .pdf):")
uploaded_file = st.file_uploader("", type=["docx", "pdf"])

if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1].lower()
    with st.spinner("Đang xử lý tệp..."):
        try:
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(uploaded_file.getbuffer())
                temp_file_path = temp_file.name

            converter = CustomDocumentConverter()
            text = converter.convert(temp_file_path, file_type=file_type)
            chunks = chunk_by_token(text, max_len=1000)
            st.success("File đã được xử lý thành công!")

            start_chunk, end_chunk = st.slider(
                "Chọn đoạn để xử lý:",
                min_value=0,
                max_value=len(chunks) - 1,
                value=(0, 5)
            )
            st.write(f"Đang xử lý từ đoạn {start_chunk + 1} đến đoạn {end_chunk + 1}...")

            selected_chunks = " ".join(chunks[start_chunk:end_chunk + 1])
            st.text_area("Nội dung các đoạn đã chọn:", selected_chunks, height=300)

            extraction_levels = {
                "Nhận biết": ExtractA,
                "Thông hiểu": ExtractB,
                "Vận dụng": ExtractC,
                "Vận dụng cao": ExtractD
            }
            selected_levels = st.multiselect(
                "Chọn mức độ câu hỏi cần trích xuất",
                options=list(extraction_levels.keys()),
                default=[]
            )

            question_counts = {}
            for level in selected_levels:
                count = st.number_input(
                    f"Số lượng câu hỏi mức '{level}'",
                    min_value=1,
                    max_value=50,
                    value=5,
                    step=1
                )
                question_counts[level] = count

            if st.button("Tạo câu hỏi", key="generate_questions"):
                if not selected_levels:
                    st.error("Vui lòng chọn ít nhất một mức độ câu hỏi.")
                elif not exam_title:
                    st.error("Vui lòng nhập tiêu đề đề kiểm tra.")
                else:
                    all_results = []
                    for level in selected_levels:
                        extractor_class = extraction_levels[level]
                        extractor = extractor_class()
                        result = extractor.run(selected_chunks, question_counts[level])
                        all_results.extend(result["quizes"])

                    st.session_state["quiz_data"] = all_results

                    for quiz in all_results:
                        st.write(f"**Câu hỏi ({quiz['level']}):** {quiz['question']}")
                        st.write("**Đáp án:**")
                        for i, option in enumerate(quiz['choices'], start=65):  # 65 là ASCII của 'A'
                            st.write(f"{chr(i)}. {option.strip()}")
                        st.write(f"**Đáp án đúng:** {quiz['answer']}")

                    file_path = export_to_word(all_results, exam_title)
                    with open(file_path, "rb") as f:
                        st.download_button(
                            label="Tải xuống file Word",
                            data=f,
                            file_name="Đề trắc nghiệm.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
        except Exception as e:
            st.error(f"Lỗi: {e}")
